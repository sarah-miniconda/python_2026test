from __future__ import annotations

import json
import re
import tkinter as tk
from dataclasses import asdict, dataclass, field
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from typing import Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ItemType = Literal["main", "sub", "cp", "exception"]

TYPE_LABELS: dict[ItemType, str] = {
    "main": "一級流程",
    "sub": "二級步驟",
    "cp": "管制點（CP）",
    "exception": "異常處理",
}

LABEL_TO_TYPE: dict[str, ItemType] = {
    label: code for code, label in TYPE_LABELS.items()
}


@dataclass
class SOPItem:
    item_type: ItemType
    number: str
    title: str
    details: list[str] = field(default_factory=list)
    owner: str = ""
    storage: str = ""

    @property
    def details_text(self) -> str:
        return "\n".join(self.details)


# =========================================================
# 文字解析
# =========================================================

MAIN_PATTERN = re.compile(
    r"^(?P<number>\d+)[.．、]\s*(?P<title>.+)$"
)

SUB_PATTERN = re.compile(
    r"^(?P<number>\d+\.\d+)[.．、]?\s*(?P<title>.+)$"
)

CP_PATTERN = re.compile(
    r"^(?P<number>CP\s*\d+)\s*[：:、.\-]?\s*(?P<title>.+)$",
    re.IGNORECASE,
)

BULLET_PATTERN = re.compile(
    r"^\s*[-–—•▪●○□☐✓✔＊*]\s*(?P<text>.+)$"
)


def normalize_line(line: str) -> str:
    return (
        line.replace("\u3000", " ")
        .replace("\t", " ")
        .strip()
    )


def parse_sop(raw_text: str) -> tuple[str, list[SOPItem]]:
    """
    規則：
    - 1. 標題：一級流程
    - 1.1 標題：二級步驟
    - CP1 標題：管制點
    - 二級步驟底下重新出現 1.、2.、3. 時，視為具體做法
    """
    lines = [
        normalize_line(line)
        for line in raw_text.splitlines()
        if normalize_line(line)
    ]

    if not lines:
        return "", []

    document_title = ""
    items: list[SOPItem] = []
    current_item: SOPItem | None = None
    current_main_number = 0

    for line in lines:
        # 先判斷二級標題。
        sub_match = SUB_PATTERN.match(line)
        if sub_match:
            current_item = SOPItem(
                item_type="sub",
                number=sub_match.group("number"),
                title=sub_match.group("title").strip(),
            )
            items.append(current_item)
            continue

        cp_match = CP_PATTERN.match(line)
        if cp_match:
            current_item = SOPItem(
                item_type="cp",
                number=cp_match.group("number").upper().replace(" ", ""),
                title=cp_match.group("title").strip(),
            )
            items.append(current_item)
            continue

        main_match = MAIN_PATTERN.match(line)
        if main_match:
            number = int(main_match.group("number"))
            title = main_match.group("title").strip()

            # 已經位於二級步驟時，1.、2.、3. 通常是操作細項。
            # 只有編號等於下一個主流程號碼，才視為新的一級流程。
            if (
                current_item is not None
                and current_item.item_type == "sub"
                and number != current_main_number + 1
            ):
                current_item.details.append(
                    f"{number}. {title}"
                )
                continue

            # 若目前是一級流程，重複從 1 開始也多半是細項。
            if (
                current_item is not None
                and current_item.item_type in {"main", "exception"}
                and current_main_number > 0
                and number <= current_main_number
            ):
                current_item.details.append(
                    f"{number}. {title}"
                )
                continue

            item_type: ItemType = (
                "exception" if "異常" in title else "main"
            )

            current_item = SOPItem(
                item_type=item_type,
                number=str(number),
                title=title,
            )
            items.append(current_item)
            current_main_number = number
            continue

        bullet_match = BULLET_PATTERN.match(line)
        if bullet_match and current_item:
            current_item.details.append(
                bullet_match.group("text").strip()
            )
            continue

        # 第一個非編號文字視為 SOP 名稱。
        if not items and not document_title:
            document_title = line
            continue

        if current_item:
            current_item.details.append(line)
        elif not document_title:
            document_title = line

    return document_title, items


# =========================================================
# Word 輸出
# =========================================================

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 9.5,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    lines = text.splitlines() or [""]

    for index, line in enumerate(lines):
        paragraph = (
            cell.paragraphs[0]
            if index == 0
            else cell.add_paragraph()
        )
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05

        run = paragraph.add_run(line)
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Microsoft JhengHei",
        )
        run.font.size = Pt(size)
        run.bold = bold

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def export_word(
    document_title: str,
    items: list[SOPItem],
    output_path: Path,
) -> None:
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft JhengHei"
    normal_style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Microsoft JhengHei",
    )
    normal_style.font.size = Pt(9.5)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(6)

    title_run = title_paragraph.add_run(
        document_title or "SOP 作業流程總表（草案）"
    )
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Microsoft JhengHei"
    title_run.font.color.rgb = RGBColor(92, 74, 104)
    title_run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Microsoft JhengHei",
    )

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    headers = [
        "流程",
        "說明",
        "具體做法",
        "承辦人",
        "資料存放",
    ]
    widths = [
        Cm(4.2),
        Cm(5.0),
        Cm(10.0),
        Cm(3.3),
        Cm(4.5),
    ]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        set_cell_text(
            cell,
            header,
            bold=True,
            size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_shading(cell, "E8DFF0")

    main_colors = {
        "main": "E4F1E8",
        "sub": "FFF6E9",
        "cp": "FFF0C9",
        "exception": "FADCD9",
    }

    current_main_title = ""

    for item in items:
        row = table.add_row()
        cells = row.cells

        for index, cell in enumerate(cells):
            cell.width = widths[index]

        if item.item_type in {"main", "exception"}:
            current_main_title = (
                f"{item.number}. {item.title}"
            )
            flow_text = current_main_title
            description_text = ""

        elif item.item_type == "sub":
            flow_text = current_main_title
            description_text = (
                f"{item.number} {item.title}"
            )

        else:
            flow_text = "重要管制點"
            description_text = (
                f"{item.number} {item.title}"
            )

        details_text = "\n".join(item.details)

        values = [
            flow_text,
            description_text,
            details_text,
            item.owner,
            item.storage,
        ]

        for index, value in enumerate(values):
            align = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index in {0, 3}
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_cell_text(
                cells[index],
                value,
                bold=(index == 0),
                align=align,
            )
            set_cell_shading(
                cells[index],
                main_colors[item.item_type],
            )

    doc.save(output_path)


# =========================================================
# GUI
# =========================================================

class SOPBuilderApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("SOP Builder v0.3")
        self.root.geometry("1480x850")
        self.root.minsize(1160, 700)

        self.items: list[SOPItem] = []
        self.selected_index: int | None = None
        self.current_project_path: Path | None = None

        self.colors = {
            "window": "#F8F4F7",
            "panel": "#FFF9FB",
            "input": "#FFFCFD",
            "green": "#E4F1E8",
            "purple": "#E8DFF0",
            "pink": "#FADCD9",
            "cream": "#FFF6E9",
            "text": "#514A54",
            "accent": "#9278A4",
            "selected": "#CDB7D8",
        }

        self.root.configure(bg=self.colors["window"])

        self._configure_style()
        self._build_ui()
        self._install_keyboard_shortcuts()
        self._install_context_menu()

    # -----------------------------------------------------
    # 外觀
    # -----------------------------------------------------

    def _configure_style(self) -> None:
        style = ttk.Style()

        if "clam" in style.theme_names():
            style.theme_use("clam")

        style.configure(
            ".",
            font=("PingFang TC", 10),
            background=self.colors["window"],
            foreground=self.colors["text"],
        )

        style.configure(
            "Title.TLabel",
            font=("PingFang TC", 20, "bold"),
            background=self.colors["window"],
            foreground="#5D5062",
        )

        style.configure(
            "Section.TLabel",
            font=("PingFang TC", 11, "bold"),
            background=self.colors["panel"],
        )

        style.configure(
            "Pastel.TLabelframe",
            background=self.colors["panel"],
            bordercolor="#DCCFDC",
        )

        style.configure(
            "Pastel.TLabelframe.Label",
            font=("PingFang TC", 11, "bold"),
            background=self.colors["panel"],
            foreground="#625469",
        )

        style.configure(
            "Pastel.TButton",
            padding=(10, 6),
            background=self.colors["purple"],
            foreground=self.colors["text"],
            bordercolor="#CDBED5",
        )

        style.map(
            "Pastel.TButton",
            background=[
                ("active", self.colors["pink"]),
                ("pressed", self.colors["selected"]),
            ],
        )

        style.configure(
            "Treeview",
            font=("PingFang TC", 10),
            rowheight=29,
            background="#FFFCFD",
            fieldbackground="#FFFCFD",
            foreground=self.colors["text"],
        )

        style.configure(
            "Treeview.Heading",
            font=("PingFang TC", 10, "bold"),
            background=self.colors["purple"],
            foreground="#4D4352",
        )

        style.map(
            "Treeview",
            background=[
                ("selected", self.colors["selected"]),
            ],
            foreground=[
                ("selected", "#3F3444"),
            ],
        )

    def _build_ui(self) -> None:
        outer = ttk.Frame(
            self.root,
            padding=14,
        )
        outer.pack(fill="both", expand=True)

        title = ttk.Label(
            outer,
            text="SOP Builder",
            style="Title.TLabel",
        )
        title.pack(anchor="w")

        subtitle = ttk.Label(
            outer,
            text="貼上原始內容、解析階層，並在輸出前人工確認與修改。",
        )
        subtitle.pack(anchor="w", pady=(0, 10))

        top_toolbar = ttk.Frame(outer)
        top_toolbar.pack(fill="x", pady=(0, 9))

        ttk.Button(
            top_toolbar,
            text="開啟專案",
            style="Pastel.TButton",
            command=self.open_project,
        ).pack(side="left")

        ttk.Button(
            top_toolbar,
            text="儲存專案",
            style="Pastel.TButton",
            command=self.save_project,
        ).pack(side="left", padx=(6, 0))

        ttk.Button(
            top_toolbar,
            text="另存專案",
            style="Pastel.TButton",
            command=lambda: self.save_project(save_as=True),
        ).pack(side="left", padx=(6, 0))

        ttk.Button(
            top_toolbar,
            text="輸出 Word",
            style="Pastel.TButton",
            command=self.export_word_file,
        ).pack(side="left", padx=(16, 0))

        title_frame = ttk.Frame(outer)
        title_frame.pack(fill="x", pady=(0, 10))

        ttk.Label(
            title_frame,
            text="SOP 名稱：",
            style="Section.TLabel",
        ).pack(side="left")

        self.document_title_var = tk.StringVar()
        ttk.Entry(
            title_frame,
            textvariable=self.document_title_var,
        ).pack(
            side="left",
            fill="x",
            expand=True,
            padx=(8, 0),
        )

        main_paned = ttk.Panedwindow(
            outer,
            orient="horizontal",
        )
        main_paned.pack(fill="both", expand=True)

        self.input_frame = self._build_input_panel(main_paned)
        self.tree_frame = self._build_tree_panel(main_paned)
        self.editor_frame = self._build_editor_panel(main_paned)

        main_paned.add(self.input_frame, weight=3)
        main_paned.add(self.tree_frame, weight=4)
        main_paned.add(self.editor_frame, weight=3)

        self.status_var = tk.StringVar(
            value="請貼上 SOP 文字，再按「解析文字」。"
        )
        ttk.Label(
            outer,
            textvariable=self.status_var,
            anchor="w",
        ).pack(fill="x", pady=(8, 0))

    def _build_input_panel(
        self,
        parent: ttk.Panedwindow,
    ) -> ttk.LabelFrame:
        frame = ttk.LabelFrame(
            parent,
            text="① 原始文字",
            padding=10,
            style="Pastel.TLabelframe",
        )

        toolbar = ttk.Frame(frame)
        toolbar.pack(fill="x", pady=(0, 8))

        ttk.Button(
            toolbar,
            text="解析文字",
            style="Pastel.TButton",
            command=self.parse_input,
        ).pack(side="left")

        ttk.Button(
            toolbar,
            text="清空",
            style="Pastel.TButton",
            command=self.clear_all,
        ).pack(side="left", padx=(6, 0))

        text_frame = ttk.Frame(frame)
        text_frame.pack(fill="both", expand=True)

        self.input_text = tk.Text(
            text_frame,
            wrap="word",
            undo=True,
            font=("PingFang TC", 11),
            bg=self.colors["input"],
            fg=self.colors["text"],
            insertbackground=self.colors["text"],
            selectbackground=self.colors["selected"],
            padx=10,
            pady=10,
            relief="flat",
            highlightthickness=1,
            highlightbackground="#DED2DD",
        )

        scrollbar = ttk.Scrollbar(
            text_frame,
            orient="vertical",
            command=self.input_text.yview,
        )
        self.input_text.configure(
            yscrollcommand=scrollbar.set
        )

        self.input_text.pack(
            side="left",
            fill="both",
            expand=True,
        )
        scrollbar.pack(side="right", fill="y")

        return frame

    def _build_tree_panel(
        self,
        parent: ttk.Panedwindow,
    ) -> ttk.LabelFrame:
        frame = ttk.LabelFrame(
            parent,
            text="② 結構化內容",
            padding=10,
            style="Pastel.TLabelframe",
        )

        toolbar = ttk.Frame(frame)
        toolbar.pack(fill="x", pady=(0, 8))

        buttons = [
            ("新增一級", lambda: self.add_item("main")),
            ("新增二級", lambda: self.add_item("sub")),
            ("新增 CP", lambda: self.add_item("cp")),
            ("刪除", self.delete_selected),
            ("上移", lambda: self.move_selected(-1)),
            ("下移", lambda: self.move_selected(1)),
        ]

        for index, (text, command) in enumerate(buttons):
            ttk.Button(
                toolbar,
                text=text,
                style="Pastel.TButton",
                command=command,
            ).pack(
                side="left",
                padx=(0 if index == 0 else 5, 0),
            )

        container = ttk.Frame(frame)
        container.pack(fill="both", expand=True)

        self.tree = ttk.Treeview(
            container,
            columns=("type", "number", "title", "owner"),
            show="headings",
            selectmode="browse",
        )

        for column, title in [
            ("type", "類型"),
            ("number", "編號"),
            ("title", "標題"),
            ("owner", "承辦人"),
        ]:
            self.tree.heading(column, text=title)

        self.tree.column(
            "type",
            width=95,
            anchor="center",
            stretch=False,
        )
        self.tree.column(
            "number",
            width=70,
            anchor="center",
            stretch=False,
        )
        self.tree.column(
            "title",
            width=260,
        )
        self.tree.column(
            "owner",
            width=100,
            anchor="center",
        )

        y_scroll = ttk.Scrollbar(
            container,
            orient="vertical",
            command=self.tree.yview,
        )
        x_scroll = ttk.Scrollbar(
            container,
            orient="horizontal",
            command=self.tree.xview,
        )

        self.tree.configure(
            yscrollcommand=y_scroll.set,
            xscrollcommand=x_scroll.set,
        )

        self.tree.grid(row=0, column=0, sticky="nsew")
        y_scroll.grid(row=0, column=1, sticky="ns")
        x_scroll.grid(row=1, column=0, sticky="ew")

        container.rowconfigure(0, weight=1)
        container.columnconfigure(0, weight=1)

        self.tree.bind(
            "<<TreeviewSelect>>",
            self.on_tree_select,
        )

        return frame

    def _build_editor_panel(
        self,
        parent: ttk.Panedwindow,
    ) -> ttk.LabelFrame:
        frame = ttk.LabelFrame(
            parent,
            text="③ 編輯選取項目",
            padding=12,
            style="Pastel.TLabelframe",
        )

        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(4, weight=1)

        ttk.Label(frame, text="類型").grid(
            row=0, column=0, sticky="w", pady=5
        )

        self.type_var = tk.StringVar(
            value=TYPE_LABELS["main"]
        )
        self.type_combo = ttk.Combobox(
            frame,
            textvariable=self.type_var,
            state="readonly",
            values=list(TYPE_LABELS.values()),
        )
        self.type_combo.grid(
            row=0, column=1, sticky="ew", pady=5
        )

        self.number_var = tk.StringVar()
        self.title_var = tk.StringVar()
        self.owner_var = tk.StringVar()
        self.storage_var = tk.StringVar()

        fields = [
            ("編號", self.number_var, 1),
            ("標題", self.title_var, 2),
            ("承辦人", self.owner_var, 3),
        ]

        for label, variable, row in fields:
            ttk.Label(frame, text=label).grid(
                row=row,
                column=0,
                sticky="w",
                pady=5,
            )
            ttk.Entry(
                frame,
                textvariable=variable,
            ).grid(
                row=row,
                column=1,
                sticky="ew",
                pady=5,
            )

        ttk.Label(frame, text="具體做法").grid(
            row=4,
            column=0,
            sticky="nw",
            pady=5,
        )

        details_container = ttk.Frame(frame)
        details_container.grid(
            row=4,
            column=1,
            sticky="nsew",
            pady=5,
        )
        details_container.rowconfigure(0, weight=1)
        details_container.columnconfigure(0, weight=1)

        self.details_text = tk.Text(
            details_container,
            wrap="word",
            undo=True,
            font=("PingFang TC", 10),
            bg=self.colors["input"],
            fg=self.colors["text"],
            insertbackground=self.colors["text"],
            selectbackground=self.colors["selected"],
            padx=8,
            pady=8,
            relief="flat",
            highlightthickness=1,
            highlightbackground="#DED2DD",
        )

        detail_scroll = ttk.Scrollbar(
            details_container,
            orient="vertical",
            command=self.details_text.yview,
        )
        self.details_text.configure(
            yscrollcommand=detail_scroll.set
        )

        self.details_text.grid(
            row=0,
            column=0,
            sticky="nsew",
        )
        detail_scroll.grid(
            row=0,
            column=1,
            sticky="ns",
        )

        ttk.Label(frame, text="資料存放").grid(
            row=5,
            column=0,
            sticky="w",
            pady=5,
        )
        ttk.Entry(
            frame,
            textvariable=self.storage_var,
        ).grid(
            row=5,
            column=1,
            sticky="ew",
            pady=5,
        )

        button_frame = ttk.Frame(frame)
        button_frame.grid(
            row=6,
            column=0,
            columnspan=2,
            sticky="ew",
            pady=(12, 0),
        )

        ttk.Button(
            button_frame,
            text="套用修改",
            style="Pastel.TButton",
            command=self.apply_changes,
        ).pack(side="left")

        ttk.Button(
            button_frame,
            text="取消選取",
            style="Pastel.TButton",
            command=self.clear_editor,
        ).pack(side="left", padx=(7, 0))

        return frame

    # -----------------------------------------------------
    # 複製貼上
    # -----------------------------------------------------

    def _install_keyboard_shortcuts(self) -> None:
        shortcuts = {
            "<Command-c>": "<<Copy>>",
            "<Command-v>": "<<Paste>>",
            "<Command-x>": "<<Cut>>",
            "<Command-a>": "<<SelectAll>>",
            "<Control-c>": "<<Copy>>",
            "<Control-v>": "<<Paste>>",
            "<Control-x>": "<<Cut>>",
            "<Control-a>": "<<SelectAll>>",
        }

        for sequence, virtual_event in shortcuts.items():
            self.root.bind_all(
                sequence,
                lambda event, action=virtual_event:
                self._generate_text_event(event, action),
            )

    @staticmethod
    def _generate_text_event(
        event: tk.Event,
        action: str,
    ) -> str:
        widget = event.widget

        if isinstance(widget, (tk.Text, tk.Entry, ttk.Entry)):
            widget.event_generate(action)
            return "break"

        return ""

    def _install_context_menu(self) -> None:
        self.context_menu = tk.Menu(
            self.root,
            tearoff=False,
        )
        self.context_menu.add_command(
            label="剪下",
            command=lambda:
            self._context_action("<<Cut>>"),
        )
        self.context_menu.add_command(
            label="複製",
            command=lambda:
            self._context_action("<<Copy>>"),
        )
        self.context_menu.add_command(
            label="貼上",
            command=lambda:
            self._context_action("<<Paste>>"),
        )
        self.context_menu.add_separator()
        self.context_menu.add_command(
            label="全選",
            command=lambda:
            self._context_action("<<SelectAll>>"),
        )

        for widget in [
            self.input_text,
            self.details_text,
        ]:
            widget.bind(
                "<Button-2>",
                self._show_context_menu,
            )
            widget.bind(
                "<Button-3>",
                self._show_context_menu,
            )

    def _show_context_menu(self, event: tk.Event) -> str:
        event.widget.focus_set()
        self.context_menu.tk_popup(
            event.x_root,
            event.y_root,
        )
        return "break"

    def _context_action(self, action: str) -> None:
        widget = self.root.focus_get()
        if widget:
            widget.event_generate(action)

    # -----------------------------------------------------
    # 解析
    # -----------------------------------------------------

    def parse_input(self) -> None:
        raw_text = self.input_text.get(
            "1.0",
            "end",
        ).strip()

        if not raw_text:
            messagebox.showwarning(
                "尚未輸入內容",
                "請先貼上 SOP 文字。",
            )
            return

        title, items = parse_sop(raw_text)

        if not items:
            messagebox.showwarning(
                "無法解析",
                "沒有辨識到一級標題、二級標題或管制點。",
            )
            return

        self.document_title_var.set(title)
        self.items = items
        self.current_project_path = None
        self.clear_editor()
        self.refresh_tree()

        self.status_var.set(
            f"解析完成，共 {len(items)} 個項目。"
        )

    def refresh_tree(
        self,
        selected_index: int | None = None,
    ) -> None:
        for item_id in self.tree.get_children():
            self.tree.delete(item_id)

        for index, item in enumerate(self.items):
            title = item.title
            if item.item_type == "sub":
                title = f"　└ {title}"

            self.tree.insert(
                "",
                "end",
                iid=str(index),
                values=(
                    TYPE_LABELS[item.item_type],
                    item.number,
                    title,
                    item.owner,
                ),
                tags=(item.item_type,),
            )

        self.tree.tag_configure(
            "main",
            background=self.colors["green"],
        )
        self.tree.tag_configure(
            "sub",
            background=self.colors["cream"],
        )
        self.tree.tag_configure(
            "cp",
            background="#FFF0C9",
        )
        self.tree.tag_configure(
            "exception",
            background=self.colors["pink"],
        )

        if (
            selected_index is not None
            and 0 <= selected_index < len(self.items)
        ):
            iid = str(selected_index)
            self.tree.selection_set(iid)
            self.tree.focus(iid)
            self.tree.see(iid)

    # -----------------------------------------------------
    # 編輯
    # -----------------------------------------------------

    def on_tree_select(
        self,
        _event: tk.Event | None = None,
    ) -> None:
        selection = self.tree.selection()
        if not selection:
            return

        index = int(selection[0])
        if not 0 <= index < len(self.items):
            return

        self.selected_index = index
        item = self.items[index]

        self.type_var.set(TYPE_LABELS[item.item_type])
        self.number_var.set(item.number)
        self.title_var.set(item.title)
        self.owner_var.set(item.owner)
        self.storage_var.set(item.storage)

        self.details_text.delete("1.0", "end")
        self.details_text.insert(
            "1.0",
            item.details_text,
        )

    def apply_changes(self) -> None:
        if self.selected_index is None:
            messagebox.showwarning(
                "尚未選取",
                "請先選取一個項目。",
            )
            return

        title = self.title_var.get().strip()
        if not title:
            messagebox.showwarning(
                "標題不可空白",
                "請填寫標題。",
            )
            return

        item = self.items[self.selected_index]
        item.item_type = LABEL_TO_TYPE[
            self.type_var.get()
        ]
        item.number = self.number_var.get().strip()
        item.title = title
        item.owner = self.owner_var.get().strip()
        item.storage = self.storage_var.get().strip()
        item.details = [
            line.strip()
            for line in self.details_text.get(
                "1.0",
                "end",
            ).splitlines()
            if line.strip()
        ]

        index = self.selected_index
        self.refresh_tree(index)
        self.status_var.set("修改已套用。")

    def clear_editor(self) -> None:
        self.selected_index = None
        self.type_var.set(TYPE_LABELS["main"])
        self.number_var.set("")
        self.title_var.set("")
        self.owner_var.set("")
        self.storage_var.set("")
        self.details_text.delete("1.0", "end")

        for selected in self.tree.selection():
            self.tree.selection_remove(selected)

    # -----------------------------------------------------
    # 新增、刪除、排序
    # -----------------------------------------------------

    def add_item(self, item_type: ItemType) -> None:
        item = SOPItem(
            item_type=item_type,
            number=self._suggest_number(item_type),
            title="新項目",
        )

        insert_index = (
            len(self.items)
            if self.selected_index is None
            else self.selected_index + 1
        )

        self.items.insert(insert_index, item)
        self.refresh_tree(insert_index)
        self.on_tree_select()

    def _suggest_number(
        self,
        item_type: ItemType,
    ) -> str:
        if item_type in {"main", "exception"}:
            numbers = [
                int(item.number)
                for item in self.items
                if (
                    item.item_type in {"main", "exception"}
                    and item.number.isdigit()
                )
            ]
            return str(max(numbers, default=0) + 1)

        if item_type == "cp":
            numbers = []
            for item in self.items:
                if item.item_type == "cp":
                    match = re.search(r"\d+", item.number)
                    if match:
                        numbers.append(int(match.group()))
            return f"CP{max(numbers, default=0) + 1}"

        parent = "1"
        start = (
            self.selected_index
            if self.selected_index is not None
            else len(self.items) - 1
        )

        for index in range(start, -1, -1):
            item = self.items[index]
            if item.item_type in {"main", "exception"}:
                parent = item.number
                break

        sub_numbers = []
        for item in self.items:
            match = re.fullmatch(
                rf"{re.escape(parent)}\.(\d+)",
                item.number,
            )
            if item.item_type == "sub" and match:
                sub_numbers.append(int(match.group(1)))

        return f"{parent}.{max(sub_numbers, default=0) + 1}"

    def delete_selected(self) -> None:
        if self.selected_index is None:
            messagebox.showwarning(
                "尚未選取",
                "請先選取要刪除的項目。",
            )
            return

        item = self.items[self.selected_index]

        if not messagebox.askyesno(
            "確認刪除",
            f"確定刪除「{item.number} {item.title}」嗎？",
        ):
            return

        index = self.selected_index
        del self.items[index]
        self.clear_editor()

        next_index = (
            min(index, len(self.items) - 1)
            if self.items
            else None
        )
        self.refresh_tree(next_index)

    def move_selected(self, direction: int) -> None:
        if self.selected_index is None:
            messagebox.showwarning(
                "尚未選取",
                "請先選取要移動的項目。",
            )
            return

        old = self.selected_index
        new = old + direction

        if not 0 <= new < len(self.items):
            return

        self.items[old], self.items[new] = (
            self.items[new],
            self.items[old],
        )

        self.selected_index = new
        self.refresh_tree(new)
        self.on_tree_select()

    # -----------------------------------------------------
    # 專案儲存
    # -----------------------------------------------------

    def save_project(
        self,
        save_as: bool = False,
    ) -> None:
        if not self.items:
            messagebox.showwarning(
                "沒有內容",
                "目前沒有可儲存的 SOP 項目。",
            )
            return

        path = self.current_project_path

        if save_as or path is None:
            selected = filedialog.asksaveasfilename(
                title="儲存 SOP 專案",
                defaultextension=".json",
                filetypes=[
                    ("SOP 專案", "*.json"),
                    ("所有檔案", "*.*"),
                ],
            )
            if not selected:
                return
            path = Path(selected)

        data = {
            "version": "0.3",
            "title": self.document_title_var.get().strip(),
            "raw_text": self.input_text.get(
                "1.0",
                "end",
            ).strip(),
            "items": [
                asdict(item)
                for item in self.items
            ],
        }

        try:
            path.write_text(
                json.dumps(
                    data,
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
        except OSError as exc:
            messagebox.showerror(
                "儲存失敗",
                str(exc),
            )
            return

        self.current_project_path = path
        self.status_var.set(f"已儲存：{path.name}")

    def open_project(self) -> None:
        selected = filedialog.askopenfilename(
            title="開啟 SOP 專案",
            filetypes=[
                ("SOP 專案", "*.json"),
                ("所有檔案", "*.*"),
            ],
        )
        if not selected:
            return

        path = Path(selected)

        try:
            data = json.loads(
                path.read_text(encoding="utf-8")
            )

            loaded_items = [
                SOPItem(**item)
                for item in data.get("items", [])
            ]
        except (
            OSError,
            json.JSONDecodeError,
            TypeError,
        ) as exc:
            messagebox.showerror(
                "開啟失敗",
                str(exc),
            )
            return

        self.document_title_var.set(
            data.get("title", "")
        )

        self.input_text.delete("1.0", "end")
        self.input_text.insert(
            "1.0",
            data.get("raw_text", ""),
        )

        self.items = loaded_items
        self.current_project_path = path
        self.clear_editor()
        self.refresh_tree()

        self.status_var.set(f"已開啟：{path.name}")

    # -----------------------------------------------------
    # Word 輸出
    # -----------------------------------------------------

    def export_word_file(self) -> None:
        if not self.items:
            messagebox.showwarning(
                "沒有內容",
                "請先解析或開啟一份 SOP。",
            )
            return

        selected = filedialog.asksaveasfilename(
            title="輸出 Word",
            defaultextension=".docx",
            initialfile=(
                self.document_title_var.get().strip()
                or "SOP作業流程總表"
            ),
            filetypes=[
                ("Word 文件", "*.docx"),
            ],
        )

        if not selected:
            return

        try:
            export_word(
                self.document_title_var.get().strip(),
                self.items,
                Path(selected),
            )
        except Exception as exc:
            messagebox.showerror(
                "輸出失敗",
                str(exc),
            )
            return

        messagebox.showinfo(
            "輸出完成",
            f"Word 已儲存至：\n{selected}",
        )

    # -----------------------------------------------------
    # 清空
    # -----------------------------------------------------

    def clear_all(self) -> None:
        if (
            self.input_text.get("1.0", "end").strip()
            or self.items
        ):
            if not messagebox.askyesno(
                "確認清空",
                "確定清空目前所有內容嗎？",
            ):
                return

        self.input_text.delete("1.0", "end")
        self.document_title_var.set("")
        self.items.clear()
        self.current_project_path = None
        self.clear_editor()
        self.refresh_tree()

        self.status_var.set("內容已清空。")


def main() -> None:
    root = tk.Tk()
    SOPBuilderApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()