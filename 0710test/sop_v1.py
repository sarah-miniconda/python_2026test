from pathlib import Path
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def is_main_heading(line: str) -> bool:
    return bool(re.match(r"^\d{1,2}[.．、]\s+", line.strip()))


def is_sub_heading(line: str) -> bool:
    return bool(re.match(r"^\d{1,2}\.\d+\s+", line.strip()))


def is_cp(line: str) -> bool:
    return bool(re.match(r"^CP\d+", line.strip(), re.I))


def clean_heading(line: str) -> str:
    return re.sub(r"^\d{1,2}[.．、]\s+", "", line.strip())


def parse_sop(raw_text: str) -> list[dict]:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    steps = []

    current = None

    for line in lines:
        if is_main_heading(line):
            if current:
                steps.append(current)

            current = {
                "type": "process",
                "title": clean_heading(line),
                "details": [],
                "owner": "待確認",
            }

            if "異常" in line:
                current["type"] = "exception"

        elif is_cp(line):
            if current:
                steps.append(current)
                current = None

            steps.append({
                "type": "decision",
                "title": line + "？",
                "details": ["重要管制點"],
                "owner": "待確認",
            })

        else:
            if current:
                current["details"].append(line)

    if current:
        steps.append(current)

    return steps


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text: str, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    lines = text.split("\n") if text else [""]

    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(size)
        run.bold = bold

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def create_word(raw_text: str, output_path: Path):
    steps = parse_sop(raw_text)

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21)
    sec.top_margin = Cm(0.9)
    sec.bottom_margin = Cm(0.9)
    sec.left_margin = Cm(0.9)
    sec.right_margin = Cm(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SOP 作業流程圖（草案）")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["流程", "說明", "承辦人"]
    for i, h in enumerate(headers):
        set_text(table.rows[0].cells[i], h, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[i], "D9EAF7")

    rows = [
        ("start", "〔開始〕", "收到作業需求", "申請者／承辦人"),
        ("arrow", "↓", "", ""),
    ]

    for step in steps:
        if step["type"] == "decision":
            flow = "◇ " + step["title"]
        elif step["type"] == "exception":
            flow = "▭ " + step["title"]
        else:
            flow = "▭ " + step["title"]

        rows.append((
            step["type"],
            flow,
            "\n".join(step["details"]),
            step["owner"],
        ))
        rows.append(("arrow", "↓", "", ""))

    rows.append(("end", "〔結束〕", "完成作業", "承辦人"))

    colors = {
        "start": "E2F0D9",
        "end": "E2F0D9",
        "process": "DDEBF7",
        "decision": "FFF2CC",
        "exception": "FCE4D6",
        "arrow": "FFFFFF",
    }

    for kind, flow, details, owner in rows:
        cells = table.add_row().cells

        if kind == "arrow":
            set_text(cells[0], flow, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_text(cells[1], "")
            set_text(cells[2], "")
        else:
            set_text(cells[0], flow, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_text(cells[1], details, size=9.5)
            set_text(cells[2], owner, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

        for cell in cells:
            shade(cell, colors.get(kind, "FFFFFF"))

    doc.save(output_path)


class SOPBuilderApp:
    def __init__(self, root):
        self.root = root
        self.root.title("SOP Builder v0.1")
        self.root.geometry("1100x700")

        self.input_box = tk.Text(root, wrap="word", font=("Microsoft JhengHei", 11))
        self.input_box.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        right = tk.Frame(root)
        right.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        self.preview_box = tk.Text(right, wrap="word", font=("Microsoft JhengHei", 11))
        self.preview_box.pack(fill="both", expand=True)

        btn_frame = tk.Frame(right)
        btn_frame.pack(fill="x", pady=8)

        tk.Button(btn_frame, text="整理預覽", command=self.preview).pack(side="left", padx=4)
        tk.Button(btn_frame, text="產生 Word", command=self.export_word).pack(side="left", padx=4)
        tk.Button(btn_frame, text="清空", command=self.clear).pack(side="left", padx=4)

    def preview(self):
        raw = self.input_box.get("1.0", "end").strip()
        steps = parse_sop(raw)

        self.preview_box.delete("1.0", "end")
        self.preview_box.insert("end", "〔開始〕\n↓\n")

        for step in steps:
            prefix = "◇" if step["type"] == "decision" else "▭"
            self.preview_box.insert("end", f"{prefix} {step['title']}\n")
            for d in step["details"]:
                self.preview_box.insert("end", f"  {d}\n")
            self.preview_box.insert("end", "↓\n")

        self.preview_box.insert("end", "〔結束〕")

    def export_word(self):
        raw = self.input_box.get("1.0", "end").strip()

        if not raw:
            messagebox.showwarning("提醒", "請先貼上 SOP 文字")
            return

        path = filedialog.asksaveasfilename(
            title="儲存 Word 檔",
            defaultextension=".docx",
            filetypes=[("Word 文件", "*.docx")],
        )

        if not path:
            return

        create_word(raw, Path(path))
        messagebox.showinfo("完成", f"已輸出：\n{path}")

    def clear(self):
        self.input_box.delete("1.0", "end")
        self.preview_box.delete("1.0", "end")


if __name__ == "__main__":
    root = tk.Tk()
    app = SOPBuilderApp(root)
    root.mainloop()